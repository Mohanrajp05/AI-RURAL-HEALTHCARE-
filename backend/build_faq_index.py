"""
FAQ Index Builder
=================

Reads the curated Q&A dataset (59 verified Q&A pairs about the app itself
-- what it does, how to use it, what diseases it covers, how each feature
works) and writes it out as backend/knowledge_base/faq_index.json: a flat
list of {"question": ..., "answer": ...} objects that faq_matcher.py
matches user questions against with zero LLM calls.

Source path lives in ONE place (FAQ_DOCX_PATH below) so it's easy to
update later if the dataset moves.

The .docx is a series of "<emoji> Category Heading" paragraphs, each
immediately followed by a 2-column QUESTION/ANSWER table. This walks the
document in true body order (python-docx's separate .paragraphs/.tables
lists lose the interleaving) so a table is always paired with the heading
that actually precedes it, and re-running this after the dataset is
edited (categories/questions added or changed) picks up the changes with
no code change needed -- same auto-discovery approach as
scrape_medical_kb.py's slug-driven design.

Run (from backend/):
    python build_faq_index.py
"""

import json
import os
import re

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Single source of truth for the dataset location.
FAQ_DOCX_PATH = os.path.join(
    os.path.dirname(BASE_DIR),
    "Disease dataset",
    "Dataset",
    "RuralHealthcareChatbot FAQ_Dataset.docx",
)

OUTPUT_JSON = os.path.join(BASE_DIR, "knowledge_base", "faq_index.json")


def _iter_block_items(document):
    """Yield each top-level Paragraph/Table in the document in true reading order.

    python-docx's document.paragraphs and document.tables are separate flat
    lists that drop the interleaving between them -- this walks the
    underlying XML body directly so a heading can be reliably paired with
    the table that immediately follows it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def extract_faq_rows(docx_path: str) -> list:
    """Return [{"question": ..., "answer": ...}, ...] for every Q&A table row.

    Skips any table whose header row isn't a QUESTION/ANSWER pair (e.g. the
    dataset's leading stats-summary table), and any row where a cell is
    empty or where both cells are identical (a category banner accidentally
    formatted as a table row rather than a heading paragraph).
    """
    document = docx.Document(docx_path)
    rows = []
    for item in _iter_block_items(document):
        if isinstance(item, Table):
            if not item.rows:
                continue
            header_cells = [c.text.strip().upper() for c in item.rows[0].cells]
            if not header_cells or "QUESTION" not in header_cells[0]:
                continue
            for row in item.rows[1:]:
                cells = row.cells
                if len(cells) < 2:
                    continue
                question = cells[0].text.strip()
                answer = cells[1].text.strip()
                if not question or not answer:
                    continue
                if question == answer:
                    continue
                rows.append({"question": question, "answer": answer})
    return rows


def main():
    if not os.path.exists(FAQ_DOCX_PATH):
        raise SystemExit(f"FAQ dataset not found at {FAQ_DOCX_PATH}")

    rows = extract_faq_rows(FAQ_DOCX_PATH)
    if not rows:
        raise SystemExit("No Q&A rows found -- dataset format may have changed.")

    os.makedirs(os.path.dirname(OUTPUT_JSON), exist_ok=True)
    with open(OUTPUT_JSON, "w", encoding="utf-8") as fh:
        json.dump(rows, fh, ensure_ascii=False, indent=2)

    print(f"Extracted {len(rows)} Q&A pairs.")
    print(f"Saved to: {OUTPUT_JSON}")


if __name__ == "__main__":
    main()
