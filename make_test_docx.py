"""Create a ~2500-word test .docx with distinctive per-section content."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document

doc = Document()
doc.add_heading("Rural Health Camp Field Report 2026", level=1)
doc.add_paragraph(
    "This report documents the annual rural health outreach camp organized by the "
    "District Health Office in collaboration with local primary health centers. The camp "
    "covered 14 villages across three taluks and ran for nine consecutive days in August 2026."
)

sections = {
    "Vaccination Drive": [
        "The vaccination drive immunized 1,847 children under five years of age against measles, "
        "polio, diphtheria, and tetanus. Cold-chain equipment was audited daily and vaccine vials "
        "remained within the two-to-eight degree Celsius range throughout the camp.",
        "Community health workers maintained a digital registry on tablet devices, and each child "
        "received a paper immunization card in addition to the electronic record. Follow-up "
        "reminders for the second dose of the measles vaccine were scheduled for January 2027.",
    ],
    "Water Sanitation Survey": [
        "A survey of 312 household drinking water sources found that 58 percent of open wells "
        "tested positive for coliform bacteria. The team distributed chlorine tablets and "
        "demonstrated safe storage practices using covered containers.",
        "Borewells serving the villages of Hosahalli and Byadagi were re-tested after treatment, "
        "and both passed the bacteriological standard for drinking water by the third week.",
    ],
    "Nutrition Program": [
        "Anganwadi centers provided fortified nutrition supplements to 962 pregnant women and "
        "mothers of infants. The program measured mid-upper arm circumference monthly and referred "
        "children with measurements below the threshold to the district nutrition unit.",
        "A kitchen-garden initiative distributed vegetable seed kits to 480 households, with "
        "demonstration plots at two schools teaching composting and seasonal crop rotation.",
    ],
    "Diabetes Screening": [
        "Random blood sugar screening covered 1,203 adults over 30 years old. Of these, 214 were "
        "found to have fasting blood glucose above 126 milligrams per deciliter and were referred "
        "to the taluk hospital for confirmatory testing and management.",
        "Dr. Lakshmi, the camp's lead physician, personally reviewed every borderline result and "
        "flagged 41 patients for urgent follow-up because of additional risk factors such as "
        "hypertension or a family history of diabetes.",
    ],
    "Eye Care Clinic": [
        "The eye care unit examined 876 patients and dispensed 502 pairs of reading glasses. "
        "Cataract referrals numbered 133, and cataract surgery camps were scheduled at the "
        "district hospital for the following month.",
        "Trained optometrists used portable autorefractors, and every patient above 45 years of "
        "age received a dilated fundus examination to check for diabetic retinopathy.",
    ],
    "Health Education Sessions": [
        "Evening health education sessions reached an estimated 3,100 villagers with topics "
        "ranging from hand hygiene to recognition of early warning signs of stroke. Sessions were "
        "conducted in Kannada and Tamil to match local language preferences.",
        "Question-and-answer sessions addressed misconceptions about antibiotics, and printed "
        "posters in both languages were distributed to every village panchayat office.",
    ],
    "Summary and Recommendations": [
        "In summary, the camp delivered 8,947 individual services across all units. The highest "
        "burden of undiagnosed disease was in diabetes and dental caries, and the review board "
        "recommends a dedicated diabetes follow-up camp in Hosahalli within six months.",
        "The review board also recommends strengthening the cold-chain logistics for the "
        "winter vaccination round and appointing two additional community health workers for the "
        "Byadagi cluster to sustain the water-sanitation gains.",
    ],
}

for title, paras in sections.items():
    doc.add_heading(title, level=2)
    for p in paras:
        doc.add_paragraph(p)

# pad with a filler section to push the doc past ~2500 words
doc.add_heading("Appendix: Field Notes", level=2)
filler = (
    "Field observations were recorded daily by the camp supervisor. Teams departed at six "
    "in the morning and returned by seven in the evening. Supplies were replenished from the "
    "district cold store twice during the camp. Attendance lists were reconciled every evening "
    "and any mismatch was investigated the same night. The mobile laboratory unit ran 1,940 "
    "hemoglobin tests and 856 malaria rapid diagnostic tests during the camp period. "
)
for i in range(14):
    doc.add_paragraph(f"Day {i+1}: {filler}")

out = r"C:\Users\Mohan Raj P\AppData\Local\Temp\opencode\rural_health_camp_report.docx"
doc.save(out)
print("saved:", out)